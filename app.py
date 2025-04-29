from flask import Flask, render_template, request, send_file, Response, jsonify
from datetime import datetime
import os
import re
import io
import docx

from dotenv import load_dotenv

from utils.confluence import extract_page_id, fetch_confluence_content, fetch_all_config_steps
from utils.pdf_generator import build_pdf
from utils.ppt_generator import build_pptx
from utils.ai_helpers import generate_summary, generate_chat_reply, generate_release_note, generate_custom_document, generate_summary_and_steps
from utils.confluence_publish import publish_release_note_to_confluence
from utils.helpers import download_image_as_base64  # 💡 NEW
from weasyprint import HTML

load_dotenv()
app = Flask("collateral_vector")


# --- Security setup ---
def check_auth(username, password):
    return username == os.getenv("APP_USERNAME") and password == os.getenv("APP_PASSWORD")

def authenticate():
    return Response(
        'Could not verify your access.\n'
        'Please provide valid credentials.', 401,
        {'WWW-Authenticate': 'Basic realm="Login Required"'}
    )

def requires_auth(f):
    from functools import wraps
    @wraps(f)
    def decorated(*args, **kwargs):
        auth = request.authorization
        if not auth or not check_auth(auth.username, auth.password):
            return authenticate()
        return f(*args, **kwargs)
    return decorated

@app.route("/", methods=["GET", "POST"])
@requires_auth
def index():
    if request.method == "POST":
        confluence_url = request.form.get("confluence_url")
        action = request.form.get("action")

        if not confluence_url:
            return render_template("index.html", message="Confluence URL is required", now=datetime.now())

        page_id = extract_page_id(confluence_url)

        if action == "pdf":
            title, content, config_steps, overview_image, all_image_filenames = fetch_confluence_content(page_id)

            ai_result = generate_summary_and_steps(content)

            ai_summary = content
            ai_steps = []

            if ai_result and "Configuration Steps:" in ai_result:
                try:
                    summary_part, steps_part = ai_result.split("Configuration Steps:")
                    ai_summary = summary_part.replace("Feature Summary:", "").strip()
                    ai_steps = [line.strip() for line in steps_part.strip().split("\n") if line.strip()]
                except Exception as e:
                    print("⚠️ Failed to parse AI output:", e)

            # 💡 Embed images as base64
            overview_image_base64 = None
            if overview_image:
                overview_image_base64 = download_image_as_base64(page_id, overview_image)

            step_images_base64 = {}
            for idx, step in enumerate(ai_steps, start=1):
                expected_filename = f"step{idx}.png"
                if expected_filename in all_image_filenames:
                    img_b64 = download_image_as_base64(page_id, expected_filename)
                    if img_b64:
                        step_images_base64[idx] = img_b64

            html_content = render_template(
                "pdf_template.html",
                title=title,
                ai_summary=ai_summary,
                ai_steps=ai_steps,
                overview_image_base64=overview_image_base64,
                step_images_base64=step_images_base64,
                now=datetime.now()
            )

            pdf_file = HTML(string=html_content).write_pdf()

            return send_file(
                io.BytesIO(pdf_file),
                as_attachment=True,
                download_name="feature_guide.pdf",
                mimetype="application/pdf"
            )

        elif action == "ppt":
            title, content, config_steps, overview_image, all_image_filenames = fetch_confluence_content(page_id)
            ppt_file = build_pptx(title, content, page_id, overview_image)
            return send_file(ppt_file, as_attachment=True, download_name="feature_guide.pptx", mimetype="application/vnd.openxmlformats-officedocument.presentationml.presentation")

        elif action == "release_note":
            title, feature_text, _, overview_image, all_image_filenames = fetch_confluence_content(page_id)
            config_steps = fetch_all_config_steps(page_id)
            release_note_text, feature_name = generate_release_note("", title, feature_text, config_steps, page_id, overview_image, all_image_filenames)
            return render_template("release_note.html", release_note=release_note_text, feature_name=feature_name)

        elif action == "custom_doc":
            return render_template("custom_doc_prompt.html", confluence_url=confluence_url)

        else:
            return render_template("index.html", message="Action not recognized.", now=datetime.now())

    return render_template("index.html", now=datetime.now())

# --- Chat endpoint ---
@app.route("/chat", methods=["POST"])
@requires_auth
def chat():
    data = request.get_json()
    user_message = data.get("user_message", "").strip()
    confluence_url = data.get("confluence_url")

    if not confluence_url:
        return jsonify({"reply": "⚠️ Missing Confluence URL."})

    page_id = extract_page_id(confluence_url)
    if not page_id:
        return jsonify({"reply": "⚠️ Could not extract page ID from URL."})

    _, page_text, _, _, _ = fetch_confluence_content(page_id)

    if not user_message or not page_text:
        return jsonify({"reply": "⚠️ Missing user message or Confluence content."})

    reply = generate_chat_reply(user_message, page_text)
    return jsonify({"reply": reply})

# --- Publish release note ---
@app.route("/publish_release_note", methods=["POST"])
@requires_auth
def publish_release_note():
    data = request.get_json()
    html_content = data.get("note_content", "")
    page_title = data.get("page_title", "New Feature Release Note")

    if not html_content:
        return jsonify({"status": "error", "message": "No release note content provided."})

    success, message = publish_release_note_to_confluence(page_title, html_content)

    if success:
        return jsonify({"status": "success", "message": message})
    else:
        return jsonify({"status": "error", "message": message})

# --- Custom document generator ---
@app.route("/generate_custom_doc", methods=["POST"])
@requires_auth
def generate_custom_doc():
    user_prompt = request.form.get("prompt")
    confluence_url = request.form.get("confluence_url")

    if not user_prompt or not confluence_url:
        return "Missing prompt or Confluence URL.", 400

    page_id = extract_page_id(confluence_url)
    _, page_text, _, _, _ = fetch_confluence_content(page_id)

    draft_text = generate_custom_document(user_prompt, page_text)

    return render_template("custom_doc_editor.html", draft_text=draft_text)
@app.route("/ask_vector", methods=["POST"])
@requires_auth
def ask_vector():
    data = request.get_json()
    question = data.get("question", "").strip()

    if not question:
        return jsonify({"reply": "⚠️ Missing question."})

    from utils.vector_index import query_index
    from utils.ai_helpers import generate_chat_reply

    chunks = query_index(question)
    if not chunks:
        return jsonify({"reply": "⚠️ No relevant content found."})

    context = "\n\n".join(chunks)
    reply = generate_chat_reply(question, context)

    return jsonify({"reply": reply})
@app.route("/index_page", methods=["POST"])
@requires_auth
def index_confluence_page():
    try:
        data = request.get_json()
        confluence_url = data.get("confluence_url")

        if not confluence_url:
            return jsonify({"error": "Missing confluence_url"}), 400

        from utils.vector_index import index_page
        index_page(confluence_url)

        return jsonify({"status": "success", "message": "Page indexed."})

    except Exception as e:
        print("❌ Error indexing page:", e)
        return jsonify({"error": str(e)}), 500


# --- Download custom document as Word ---
@app.route("/download_word", methods=["POST"])
@requires_auth
def download_word():
    data = request.get_json()
    doc_text = data.get("document_text", "")

    doc = docx.Document()

    for line in doc_text.split('\n'):
        line = line.strip()

        if not line:
            continue

        line = re.sub(r"\*\*(.*?)\*\*", r"\1", line)
        line = re.sub(r"__(.*?)__", r"\1", line)

        if line.startswith("# ") or (line.endswith(":") and not line.startswith("-")):
            clean_line = line.replace("# ", "").rstrip(":")
            p = doc.add_paragraph()
            run = p.add_run(clean_line)
            run.bold = True
            run.font.size = docx.shared.Pt(14)
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(line[2:], style='ListBullet')
        else:
            doc.add_paragraph(line)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(file_stream, as_attachment=True, download_name="custom_document.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

# --- Run locally ---
if __name__ == "__main__":
    app.run(debug=True)



# (your other imports)





